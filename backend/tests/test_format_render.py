"""生成器 v2 版式测试：合同感排版（docx/pdf）+ md 内容规范。

覆盖：正文不再用「1.1/2.1」工程编号、总价大小写并用；
docx 标题居中加粗且不用 Word 默认蓝色 Heading；正文宋体小四、
首行缩进两字符、两端对齐；docx/pdf 都含双方信息表与签署盖章区。
"""

from __future__ import annotations

import re
from pathlib import Path

import pytest

from backend.app.parser import extract_text
from backend.eval.generate_samples import (
    SPECS,
    _cn_upper_amount,
    render_contract,
    render_docx,
    render_pdf,
)


def test_md_body_no_engineer_numbering_and_amount_uppercase() -> None:
    """正文不再是 1.1/2.1 工程编号，且总价大小写并用（大写 + 小写锚点都在）。"""
    md = render_contract(SPECS[0])
    numbered = [line for line in md.splitlines() if re.match(r"^\d+\.\d+\s", line)]
    assert numbered == []
    assert "壹佰万元整" in md
    assert "1,000,000" in md


def test_cn_upper_amount_values() -> None:
    """中文大写金额转换：样本值必须转换正确（渲染进正文要能看）。"""
    assert _cn_upper_amount("1,000,000") == "壹佰万元整"
    assert _cn_upper_amount("560,000") == "伍拾陆万元整"
    assert _cn_upper_amount("800,000") == "捌拾万元整"
    assert _cn_upper_amount("2,000,000") == "贰佰万元整"


def test_docx_title_centered_bold_not_default_heading(tmp_path: Path) -> None:
    """标题居中加粗，且不用 Word 默认 Heading（避免蓝色标题）。"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    path = tmp_path / "title.docx"
    render_docx(SPECS[0], path)
    doc = Document(path)
    title = doc.paragraphs[0]
    assert title.text == SPECS[0].title
    assert title.style.name == "Normal"
    assert title.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert title.runs and title.runs[0].bold is True


def test_docx_body_chinese_typography(tmp_path: Path) -> None:
    """正文样式：Normal 宋体 12pt；正文段落首行缩进两字符 + 两端对齐。"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    path = tmp_path / "body.docx"
    render_docx(SPECS[0], path)
    doc = Document(path)
    normal = doc.styles["Normal"]
    # 东亚洲字体需从 rPr/rFonts 读（font.name 只覆盖 ascii/hAnsi）
    east = normal.element.rPr.rFonts.get(qn("w:eastAsia"))
    assert east == "宋体"
    assert normal.font.size.pt == 12

    # 第一条后的正文段落：首行缩进 > 0 且两端对齐
    body = next(p for p in doc.paragraphs if p.text.startswith("乙方向甲方供应"))
    assert body.paragraph_format.first_line_indent is not None
    assert body.paragraph_format.first_line_indent > 0
    assert body.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY


def test_docx_has_party_and_signature_furniture(tmp_path: Path) -> None:
    """docx 含双方信息表（信用代码）与签署盖章区。"""
    path = tmp_path / "furniture.docx"
    render_docx(SPECS[0], path)
    text = extract_text(path)
    assert "统一社会信用代码" in text
    assert "甲方（盖章）" in text
    assert "乙方（盖章）" in text


def test_pdf_a4_with_party_and_signature(tmp_path: Path) -> None:
    """PDF 为 A4、含双方信息与签署区，关键金额仍可抽回。"""
    from pypdf import PdfReader

    path = tmp_path / "furniture.pdf"
    render_pdf(SPECS[0], path)
    reader = PdfReader(str(path))
    page = reader.pages[0]
    # A4 宽 595pt（允许 1pt 误差）
    assert abs(float(page.mediabox.width) - 595.0) < 1.0
    text = extract_text(path)
    assert "统一社会信用代码" in text
    assert "甲方（盖章）" in text
    assert "1,000,000" in text
    # 大写短语可能被段落换行断开（PDF 按行抽取），忽略空白/换行后仍应完整可读
    compact = text.replace(" ", "").replace("\n", "")
    assert "壹佰万元整" in compact


@pytest.mark.parametrize("idx", range(len(SPECS)))
def test_all_specs_render_docx_pdf_roundtrip(tmp_path: Path, idx: int) -> None:
    """全部 spec 都能出 docx/pdf，且正文关键锚点（金额/条款）不丢。"""
    spec = SPECS[idx]
    docx_path = tmp_path / f"{idx}.docx"
    pdf_path = tmp_path / f"{idx}.pdf"
    render_docx(spec, docx_path)
    render_pdf(spec, pdf_path)
    for fmt_path in (docx_path, pdf_path):
        text = extract_text(fmt_path)
        assert spec.contract_no in text
        assert "第一条" in text
        assert spec.total_amount in text
