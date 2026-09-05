"""parser 单测：文本抽取 + 按「第X条」切分 + 超长条款续切 + 无条文结构兜底。"""

from pathlib import Path

import pytest

from backend.app.config import BASE_DIR
from backend.app.parser import Clause, chunk_for_index, extract_text, split_clauses


def _sample01_text() -> str:
    return (
        BASE_DIR
        / "data"
        / "contracts"
        / "sample_01_电子元件采购合同_正常.md"
    ).read_text(encoding="utf-8")


def test_split_clauses_on_real_sample() -> None:
    clauses = split_clauses(_sample01_text())
    refs = [c.ref for c in clauses]
    # 前言 + 第一条..第十条
    assert len(clauses) >= 11
    assert "前言" in [c.title for c in clauses]
    assert "第一条" in refs and "第十条" in refs
    # 每个条款块自含条款头（证据可回指原文）
    for c in clauses:
        if c.ref:
            assert c.text.startswith(c.ref)
    # 付款条款里应能找到金额证据
    pay = next(c for c in clauses if c.title.startswith("第二条"))
    assert "1,000,000" in pay.text


def test_split_clauses_without_marker_returns_empty() -> None:
    assert split_clauses("这是一段没有任何条款编号的普通文本。") == []


def test_split_chapter_style_like_uniform_contract() -> None:
    """章节式文本（校服合同那种「一、二、三」）也能按章切分。"""
    text = (
        "甲方（采购方）：广州市晨光实验中学\n"
        "一、校服材质、数量、单价等明细\n"
        "1、夏季运动服套装，单价 150 元；\n"
        "二、单个学生校服的总价\n"
        "单个学生校服的总价为人民币（大写）陆佰贰拾元整。\n"
        "三、质量要求\n"
        "1、全新设计制造，无次品。\n"
        "2、经多次洗擦而不褪色。\n"
    )
    clauses = split_clauses(text)
    refs = [c.ref for c in clauses]
    # 前言（甲乙信息）独立成块，其后才是章节序列
    assert refs[0] == ""
    assert refs[1:] == [
        "一、校服材质、数量、单价等明细",
        "二、单个学生校服的总价",
        "三、质量要求",
    ]
    assert clauses[0].title == "前言"
    # 章节正文里的「1、2、」子条不得被当成新章节头
    assert all(c.text.startswith(c.ref) for c in clauses)
    assert "150 元" in clauses[1].text


def test_split_prefers_tiao_when_both_styles_exist() -> None:
    """同时出现「第X条」与「一、」时，以「第X条」为顶级结构（章节行算正文）。"""
    text = (
        "第一条 总则\n"
        "1、双方按民法典订立本合同。\n"
        "第二条 其他约定\n"
        "一、双方确认本合同附件与正文具有同等效力。\n"
    )
    clauses = split_clauses(text)
    refs = [c.ref for c in clauses]
    assert refs == ["第一条", "第二条"]
    # 该行属第二条正文，不得被拆成第三个条款
    assert "同等效力" in clauses[1].text


def test_long_chapter_clause_keeps_header_on_continuation() -> None:
    """章节式超长条款续切时，续块仍带章节头（如 五、…（续））。"""
    body = "乙方应依照确认的样衣组织生产，并送有资质机构检验。\n" * 40
    text = f"五、校服的生产加工与送检\n{body}\n六、交货时间、地点及货物包装\n乙方应于 2026 年 8 月 10 日前交货。\n"
    chunks = chunk_for_index(text, max_chars=200)
    assert len(chunks) > 2
    heads = [c for c in chunks if c.ref.startswith("五、")]
    assert any("（续）" in c.title for c in heads)
    assert all(c.ref.startswith("五、") for c in heads)


def test_long_clause_keeps_header_on_continuation() -> None:
    long_body = "乙方逾期交付的，每逾期一日应向甲方支付违约金，甲方并有权顺延付款。\n" * 40
    text = f"第一条 违约责任\n{long_body}\n第二条 其他\n本合同未尽事宜双方协商。\n"
    chunks = chunk_for_index(text, max_chars=200)
    assert len(chunks) > 1  # 超长条款被续切
    assert all(len(c.text) <= 200 for c in chunks)
    assert chunks[0].text.startswith("第一条")
    # 续块带条款头且不丢条款引用
    tails = [c for c in chunks[1:] if c.text.startswith("第一条")]
    assert any("（续）" in c.text for c in tails)
    assert all(c.ref for c in chunks if c.ref)  # 续块 ref 保留


def test_fallback_generic_split_when_no_clause() -> None:
    body = "本合同没有条文结构，只是很长的一段描述。\n" * 30
    chunks = chunk_for_index(body, max_chars=150)
    assert len(chunks) > 1
    assert all(len(c.text) <= 150 for c in chunks)


def test_extract_text_txt(tmp_path: Path) -> None:
    p = tmp_path / "demo.txt"
    p.write_text("甲方：测试公司\n乙方：供应商\n", encoding="utf-8")
    assert "测试公司" in extract_text(p)


def test_extract_text_docx(tmp_path: Path) -> None:
    from docx import Document

    p = tmp_path / "demo.docx"
    doc = Document()
    doc.add_paragraph("第一条 测试条款")
    doc.add_paragraph("第二条 保密条款")
    doc.save(p)
    text = extract_text(p)
    assert "第一条 测试条款" in text
    assert "第二条 保密条款" in text


def test_extract_text_unsupported_suffix(tmp_path: Path) -> None:
    p = tmp_path / "demo.xyz"
    p.write_text("hi", encoding="utf-8")
    with pytest.raises(ValueError):
        extract_text(p)
