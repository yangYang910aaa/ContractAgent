"""真实文件格式测试：docx 表格抽取 + PDF 中文抽取 + 生成器三格式输出。"""

from pathlib import Path

import pytest

from backend.app.parser import extract_text
from backend.eval.generate_samples import SPECS, _cjk_font_path, render_docx, render_pdf


def test_docx_table_content_extracted(tmp_path: Path) -> None:
    """Word 里用表格写的付款期次不能被漏掉（parser 需按文档顺序读表格）。"""
    from docx import Document

    path = tmp_path / "table.docx"
    doc = Document()
    doc.add_paragraph("第一条 付款方式")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "预付款"
    table.cell(0, 1).text = "480,000"
    table.cell(1, 0).text = "验收后支付"
    table.cell(1, 1).text = "320,000"
    doc.save(str(path))

    text = extract_text(path)
    assert "预付款" in text
    assert "480,000" in text
    assert "验收后支付" in text


def test_pdf_roundtrip_chinese(tmp_path: Path) -> None:
    """生成的中文 PDF 应能被 pypdf 抽回关键字段（无中文字体环境则跳过）。"""
    if _cjk_font_path() is None:
        pytest.skip("本机没有可用中文字体，跳过 PDF 回环测试")
    path = tmp_path / "sample.pdf"
    render_pdf(SPECS[0], path)
    text = extract_text(path)
    assert "预付款" in text
    assert "1,000,000" in text


def test_render_docx_and_pdf_outputs(tmp_path: Path) -> None:
    """生成器对同一 spec 可同时产出 docx/pdf，内容与 md 同源。"""
    if _cjk_font_path() is None:
        pytest.skip("本机没有可用中文字体，跳过 PDF 生成测试")
    docx_path = tmp_path / "sample.docx"
    pdf_path = tmp_path / "sample.pdf"
    render_docx(SPECS[0], docx_path)
    render_pdf(SPECS[0], pdf_path)
    assert "合同编号：HT-2026-0101" in extract_text(docx_path)
    assert "合同编号：HT-2026-0101" in extract_text(pdf_path)
