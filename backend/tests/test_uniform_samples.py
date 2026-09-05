"""校服式（gov_goods）样本测试：章节式正文/三格式渲染/规则期望（v3 试点）。"""

from datetime import date
from decimal import Decimal
from pathlib import Path

import pytest

from backend.app.parser import extract_text, split_clauses
from backend.app.rules import evaluate, grade_report
from backend.app.schemas import ContractModel, Grade, PaymentTerm, Severity
from backend.eval.generate_samples import UNIFORM_SPECS, _body_for, _warranty_text, render_docx, render_pdf


def _uniform_model(spec) -> ContractModel:
    """按 spec 的期望值构造 ContractModel（等价真实 LLM 抽取 + 归一化后的结果）。"""
    return ContractModel(
        contract_kind="gov_goods",
        buyer=spec.buyer,
        supplier=spec.supplier,
        signature_date=date(2026, 3, 10),
        effective_date=date(2026, 3, 10),
        expiry_date=date(2027, 9, 30),
        total_amount=Decimal("198400"),
        currency="人民币",
        payment_schedule=[
            PaymentTerm(name="一次性付清", amount=Decimal("198400"), percent=100.0),
        ],
        penalty_rate=spec.penalty_daily_percent,
        warranty_months=spec.warranty_months,
    )


def _risk_types(model: ContractModel) -> set[str]:
    return {r.risk_type for r in evaluate(model)}


def test_warranty_text_years_and_months() -> None:
    """质保月数 → 正文写法：整年按年写（贴真实合同），缺陷按月写。"""
    assert _warranty_text(24) == "2 年"
    assert _warranty_text(6) == "6 个月"


def test_uniform_md_splits_into_14_chapters() -> None:
    """校服正文（一、二、…章节式）应被 parser 按章节切分，子条不误切。"""
    body = _body_for(UNIFORM_SPECS[0])
    clauses = split_clauses(body)
    refs = [c.ref for c in clauses]
    assert refs[0] == ""  # 前言（标题/编号/双方/鉴于段）
    assert refs[1:] == [
        "一、校服材质、数量、单价等明细",
        "二、单个学生校服的总价",
        "三、质量要求",
        "四、校服的样式与封样",
        "五、校服的生产加工与送检",
        "六、交货时间、地点及货物包装",
        "七、校服验收",
        "八、售后服务与附加服务",
        "九、付款日期与方式",
        "十、履约保证金",
        "十一、违约责任",
        "十二、合同的解除",
        "十三、争议解决",
        "十四、其他事项与附则",
    ]
    # 章节内子条（1、2、…）与明细表内容留在所在章节，不产生新的顶级块
    first = clauses[1]
    assert "夏季运动服套装" in first.text
    assert "198,400" in first.text
    assert not any(c.ref.startswith("1、") for c in clauses)


def test_normal_sample06_zero_risk_pass() -> None:
    """sample_06（质保 2 年 / 违约金 0.05%）应零风险 pass。"""
    spec = UNIFORM_SPECS[0]
    assert spec.sample_id == "sample_06"
    risks = evaluate(_uniform_model(spec))
    assert risks == []
    assert grade_report(risks) == Grade.pass_


def test_defect_sample07_hits_warranty_and_penalty() -> None:
    """sample_07（质保 6 个月 + 违约金日 1.5%）应 fail，命中 P-02 与违约金。"""
    spec = UNIFORM_SPECS[1]
    assert spec.sample_id == "sample_07"
    risks = evaluate(_uniform_model(spec))
    types = {r.risk_type for r in risks}
    assert types == {"warranty_too_short", "penalty_rate_too_high"}
    assert all(r.severity == Severity.high for r in risks)
    assert grade_report(risks) == Grade.fail
    # 质保风险引 P-02；违约金风险引 P-03（细则第三条明文 1% 上限，语料 v2 起可回指）
    assert {r.policy_ref for r in risks} == {"P-02", "P-03"}


@pytest.mark.parametrize("idx", range(len(UNIFORM_SPECS)))
def test_uniform_specs_render_docx_pdf_roundtrip(tmp_path: Path, idx: int) -> None:
    """两版校服样本都能出 md/docx/pdf，且关键锚点（编号/章节头/总价款/质保句）不丢。"""
    spec = UNIFORM_SPECS[idx]
    body = _body_for(spec)
    docx_path = tmp_path / f"{spec.sample_id}.docx"
    pdf_path = tmp_path / f"{spec.sample_id}.pdf"
    render_docx(spec, docx_path, body=body)
    render_pdf(spec, pdf_path, body=body)
    for fmt_path in (docx_path, pdf_path):
        text = extract_text(fmt_path)
        compact = text.replace(" ", "").replace("\n", "")
        assert spec.contract_no in text
        assert "一、校服材质、数量、单价等明细" in text
        assert "198,400" in text
        # 质保表述按原文（空格/换行折叠后）完整可读：s06=质保2年、s07=质保6个月
        assert _warranty_text(spec.warranty_months).replace(" ", "") in compact
    # md 正文与 docx/pdf 同源：金额/质保表述一致
    assert "198,400" in body


def test_uniform_docx_chapter_header_bold_hei(tmp_path: Path) -> None:
    """校服章节头在 docx 里应黑体加粗（格式渲染器 D7 遗留点：支持章节式头）。"""
    from docx import Document
    from docx.oxml.ns import qn

    spec = UNIFORM_SPECS[0]
    path = tmp_path / "chapter.docx"
    render_docx(spec, path, body=_body_for(spec))
    doc = Document(path)
    head = next(p for p in doc.paragraphs if p.text.startswith("一、校服材质"))
    run = head.runs[0]
    assert run.font.bold is True
    assert run._element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"


def test_uniform_docx_has_detail_and_attach_tables(tmp_path: Path) -> None:
    """校服 docx 含明细表与附件清单表（表头黑体，行内容完整）。"""
    from docx import Document

    spec = UNIFORM_SPECS[0]
    path = tmp_path / "tables.docx"
    render_docx(spec, path, body=_body_for(spec))
    doc = Document(path)
    text_tables = [[c.text for c in t.rows[0].cells] for t in doc.tables]
    assert any(cells[0] == "序号" and "金额（元）" in cells for cells in text_tables)
    assert any(cells[0] == "学年" and "单人学年总价（元）" in cells for cells in text_tables)
    detail = next(t for t in doc.tables if t.rows[0].cells[0].text == "序号")
    assert [c.text for c in detail.rows[1].cells] == ["1", "夏季运动服套装", "涤棉混纺（短袖上衣、长裤各一件）", "150", "640", "96,000"]
    # 表头黑体加粗落在 runs[0]（_fill_cell 已清空残留空 run）
    assert detail.rows[0].cells[1].paragraphs[0].runs[0].font.bold is True
