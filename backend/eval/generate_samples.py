"""第 1 版合成合同生成器。

生成 5 份中文「采购合同」样本到 data/contracts/（md / docx / pdf 三格式）:
- sample_01 / sample_02:无缺陷（正常放行）
- sample_03:违约金比例过高 + 责任上限过低 + 分项加总与总额不一致
- sample_04:预付款 60%（超 30% 政策）+ 缺失保密条款
- sample_05:质保期 6 个月(不足 12)+ 保密期 60 个月(超 36)

写法:每份合同按「第X条」成块(Phase 1 parser 将按此边界切分)，
内容由参数化 spec 渲染，同一种缺陷形态可复现、可批量扩展。
全部为合成数据，不含任何真实公司/个人信息；docx/pdf 只是"格式真实"，
用于验证 PDF/Word 上传链路，内容仍是合成的（合规红线）。

用法：
    python -m backend.eval.generate_samples
"""

from __future__ import annotations

import os
from dataclasses import dataclass, field
from pathlib import Path

OUTPUT_DIR = Path(__file__).resolve().parents[2] / "data" / "contracts"


@dataclass
class SampleSpec:
    """一份合成合同的参数化规格。"""

    sample_id: str  # 样本编号（如 sample_01，评测对齐用）
    filename: str  # 输出文件名（含缺陷特征，便于人眼区分）
    contract_no: str  # 合同编号（渲染进正文）
    title: str  # 合同标题
    buyer: str  # 甲方（采购方）名称
    supplier: str  # 乙方（供应商）名称
    signature_date: str  # 签署日期（中文文本，如 2026年3月10日）
    effective_date: str  # 生效日期
    expiry_date: str  # 到期日
    currency: str = "人民币"  # 币种
    total_amount: str = "1,000,000"  # 合同总额（元，字符串保留千分位，测抽取鲁棒性）
    # 分项：[(名称, 金额, 比例)]，用于金额一致性核验
    payment_terms: list = field(default_factory=list)  # 付款期次（名称/金额/占总额百分比）
    # 首期（预付款）比例，百分比数值
    prepayment_percent: float = 20.0  # 预付款占总额比例（%）
    warranty_months: int = 24  # 质保期（月）
    # 违约金：每日比例（百分比数值），如 0.05 表示日 0.05%
    penalty_daily_percent: float = 0.05
    # 责任上限（占合同总额百分比），None 表示未单独约定上限
    liability_cap_percent: float | None = 100.0
    confidentiality_months: int = 24  # 保密期（月）
    confidentiality_clause: bool = True  # True=渲染保密条款；False=整节缺失（构造缺陷）
    termination_notice_days: int = 30  # 提前解约通知期（天）
    ip_ownership: str = "定制成果知识产权归甲方（采购方）所有"  # IP 权属表述
    governing_law: str = "中华人民共和国法律"  # 适用法律
    note: str = ""  # 缺陷说明（写进生成清单）


SPECS: list[SampleSpec] = [
    # sample_01：正常合同（对照基线：字段齐全、金额一致）
    SampleSpec(
        sample_id="sample_01",
        filename="sample_01_电子元件采购合同_正常.md",
        contract_no="HT-2026-0101",
        title="电子元件采购合同",
        buyer="星辰智造科技有限公司",
        supplier="华芯电子有限公司",
        signature_date="2026年3月10日",
        effective_date="2026年3月10日",
        expiry_date="2027年3月9日",
        payment_terms=[
            ("预付款", "200,000", 20),
            ("验收合格后支付", "800,000", 80),
        ],
        note="正常合同：字段齐全、金额一致、条款合规。",
    ),
    # sample_02：正常合同（质保恰为下限 12 个月，验证边界合规）
    SampleSpec(
        sample_id="sample_02",
        filename="sample_02_办公设备采购合同_正常.md",
        contract_no="HT-2026-0102",
        title="办公设备采购合同",
        buyer="晨光数据服务有限公司",
        supplier="联创办公设备有限公司",
        signature_date="2026年4月2日",
        effective_date="2026年4月2日",
        expiry_date="2026年10月1日",
        total_amount="560,000",
        payment_terms=[
            ("预付款", "112,000", 20),
            ("到货验收后 30 日内", "448,000", 80),
        ],
        prepayment_percent=20.0,
        warranty_months=12,
        penalty_daily_percent=0.05,
        liability_cap_percent=100.0,
        confidentiality_months=24,
        termination_notice_days=45,
        note="正常合同：质保恰为 12 个月（政策下限，合规）。",
    ),
    # sample_03：缺陷 违约金日1.5% / 责任上限5% / 分项加总≠总额
    SampleSpec(
        sample_id="sample_03",
        filename="sample_03_服务器采购合同_违约金超限_金额不一致.md",
        contract_no="HT-2026-0103",
        title="服务器及配套软件采购合同",
        buyer="星辰智造科技有限公司",
        supplier="云启信息技术有限公司",
        signature_date="2026年5月15日",
        effective_date="2026年5月20日",
        expiry_date="2027年5月19日",
        total_amount="1,000,000",
        payment_terms=[
            ("预付款", "200,000", 20),
            ("第二批（到货后）", "500,000", 50),
            ("第三批（验收后）", "400,000", 40),
        ],
        prepayment_percent=20.0,
        warranty_months=24,
        penalty_daily_percent=1.5,  # 缺陷①：违约金日 1.5%，畸高
        liability_cap_percent=5.0,  # 缺陷②：责任上限仅 5%，过低
        confidentiality_months=24,
        note="缺陷：违约金率过高 + 责任上限过低；且分项金额 20+50+40=110 万 ≠ 总额 100 万（金额不一致）。",
    ),
    # sample_04：缺陷 预付款60%（超30%上限）/ 全篇缺保密条款
    SampleSpec(
        sample_id="sample_04",
        filename="sample_04_原材料采购合同_预付款超限_缺保密条款.md",
        contract_no="HT-2026-0104",
        title="电子原材料采购合同",
        buyer="晨光数据服务有限公司",
        supplier="宏远材料科技有限公司",
        signature_date="2026年6月1日",
        effective_date="2026年6月1日",
        expiry_date="2026年12月31日",
        total_amount="800,000",
        payment_terms=[
            ("预付款", "480,000", 60),  # 缺陷①：预付款 60%
            ("验收合格后支付", "320,000", 40),
        ],
        prepayment_percent=60.0,  # 缺陷①：超过政策上限 30%
        warranty_months=12,
        confidentiality_clause=False,  # 缺陷②：缺失保密条款
        note="缺陷：预付款比例 60% 超政策上限 30%；全篇无保密条款。",
    ),
    # sample_05：缺陷 质保6个月（<12）/ 保密期60个月（>36）
    SampleSpec(
        sample_id="sample_05",
        filename="sample_05_软件采购合同_质保过短_保密期过长.md",
        contract_no="HT-2026-0105",
        title="企业管理软件采购合同",
        buyer="星辰智造科技有限公司",
        supplier="智联软件股份有限公司",
        signature_date="2026年7月8日",
        effective_date="2026年7月8日",
        expiry_date="2027年7月7日",
        total_amount="2,000,000",
        payment_terms=[
            ("预付款", "400,000", 20),
            ("上线验收后支付", "1,600,000", 80),
        ],
        prepayment_percent=20.0,
        warranty_months=6,  # 缺陷①：质保 6 个月 < 政策下限 12
        confidentiality_months=60,  # 缺陷②：保密期 60 个月 > 政策上限 36
        note="缺陷：质保期 6 个月不足 12 个月；保密期 60 个月超过 36 个月上限。",
    ),
]


def _money(amount_str: str) -> str:
    """金额字符串后补单位，正文统一为「xxx 元」格式。"""
    return f"{amount_str} 元"


def _payment_lines(spec: SampleSpec) -> list[str]:
    """付款方式正文行（不带序号前缀，编号由 render_contract 统一生成）。"""
    lines = [f"双方约定按如下期次支付合同价款（币种：{spec.currency}）："]
    cn_ordinals = ["一", "二", "三", "四", "五"]
    for idx, (name, amount, _) in enumerate(spec.payment_terms):
        lines.append(f"（{cn_ordinals[idx]}）{name}：{_money(amount)}；")
    lines.append(f"合同总价款为 {_money(spec.total_amount)}。")
    return lines


def _penalty_lines(spec: SampleSpec) -> list[str]:
    """生成违约责任正文（违约金率 + 责任上限两句）。

    分支：配置了责任上限 → 渲染上限句；未配置 → 写「按法律规定承担」兜底句，
    避免正文出现空条款。
    """
    lines = [
        f"乙方逾期交付的，每逾期一日按合同总价款的 "
        f"{spec.penalty_daily_percent:g}% 向甲方支付违约金。"
    ]
    # 分支 1：有责任上限配置 → 正文写明上限比例
    if spec.liability_cap_percent is not None:
        lines.append(
            f"除违约金外，乙方对甲方承担的赔偿责任总额以合同总价款的 "
            f"{spec.liability_cap_percent:g}% 为上限。"
        )
    # 分支 2：未配置上限 → 写法定兜底句（正文不出现空条款）
    else:
        lines.append("双方按法律规定承担违约责任。")
    return lines


def _confidentiality_lines(spec: SampleSpec) -> list[str]:
    """生成保密条款正文；返回空列表表示该节不渲染（缺保密条款缺陷）。"""
    # 分支：confidentiality_clause=False（如 sample_04）→ 返回空，整节不渲染
    if not spec.confidentiality_clause:
        return []
    return (
        "双方对因履行本合同而知悉的对方商业秘密负有保密义务。",
        f"保密期限自本合同终止之日起 {spec.confidentiality_months} 个月。",
    )


def render_contract(spec: SampleSpec) -> str:
    """按条款动态编号渲染一份合同正文（缺保密条款时整节不渲染、条款顺延）。"""
    cn_numbers = [
        "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
        "十一", "十二", "十三", "十四", "十五",
    ]
    # (条款名, 正文行列表) —— 顺序即合同条款顺序
    sections: list[tuple[str, list[str]]] = [
        (
            "合同标的与总价款",
            [
                f"乙方向甲方供应本合同项下货物/服务，合同总价款为 "
                f"{_money(spec.total_amount)}（币种：{spec.currency}）。"
            ],
        ),
        ("付款方式", _payment_lines(spec)),
        (
            "交付与验收",
            [
                "乙方应于本合同生效后 45 日内完成交付。",
                "甲方应在收到货物后 10 个工作日内组织验收，验收合格标准以双方确认的技术规范为准。",
            ],
        ),
        (
            "质量保证",
            [
                f"乙方对所供货物提供自验收合格之日起 {spec.warranty_months} 个月的质保期。",
                "质保期内出现质量问题，乙方应在 7 日内免费维修或更换。",
            ],
        ),
        ("违约责任", _penalty_lines(spec)),
    ]
    # 分支：需要保密条款才把该节加入，否则条款顺延（贴近真实"缺失"合同）
    if spec.confidentiality_clause:
        sections.append(("保密条款", _confidentiality_lines(spec)))
    sections.extend(
        [
            (
                "知识产权",
                [
                    f"{spec.ip_ownership}。",
                    "乙方保证交付物不侵犯任何第三方知识产权，因此产生的索赔由乙方承担。",
                ],
            ),
            (
                "合同期限与终止",
                [
                    f"本合同自 {spec.signature_date} 签署，自 {spec.effective_date} 生效，"
                    f"有效期至 {spec.expiry_date}。",
                    f"任何一方提前终止本合同的，应提前 "
                    f"{spec.termination_notice_days} 日书面通知对方。",
                ],
            ),
            (
                "争议解决与适用法律",
                [
                    f"本合同适用{spec.governing_law}。",
                    "因本合同产生的争议，双方应友好协商；协商不成的，"
                    "提交甲方所在地人民法院诉讼解决。",
                ],
            ),
            (
                "其他",
                ["本合同一式两份，双方各执一份，自双方盖章之日起生效。"],
            ),
        ]
    )

    parts: list[str] = [
        f"# {spec.title}",
        "",
        f"合同编号：{spec.contract_no}",
        "",
        "甲方（采购方）：" + spec.buyer,
        "乙方（供应商）：" + spec.supplier,
        "",
    ]
    for idx, (title, lines) in enumerate(sections):
        parts.append(f"第{cn_numbers[idx]}条 {title}")
        for sub_idx, line in enumerate(lines, start=1):
            parts.append(f"{idx + 1}.{sub_idx} {line}")
        parts.append("")
    return "\n".join(parts)


def _cjk_font_path() -> Path | None:
    """找系统中文字体（PDF 嵌入用）：优先单文件 TTF，TTC 集合其次。"""
    font_dir = Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts"
    for name in ("simhei.ttf", "simkai.ttf", "msyh.ttc", "simsun.ttc"):
        path = font_dir / name
        if path.exists():
            return path
    return None


def _escape_pdf_text(text: str) -> str:
    """reportlab Paragraph 走 XML，转义 & < > 防解析报错。"""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def render_docx(spec: SampleSpec, path: Path) -> None:
    """按 spec 生成 Word 版样本（正文与 md 一致，供 .docx 上传链路测试）。"""
    from docx import Document  # 延迟导入：只在生成 Word 时拉 python-docx

    doc = Document()
    for line in render_contract(spec).splitlines():
        if not line.strip():
            continue
        # 分支：markdown 标题行 → 转成 Word 一级标题（去掉 "# " 前缀）
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        else:
            doc.add_paragraph(line)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def render_pdf(spec: SampleSpec, path: Path) -> None:
    """按 spec 生成 PDF 版样本（中文用系统字体嵌入，reportlab 自动分页）。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate

    font_path = _cjk_font_path()
    # 分支：找不到中文字体 → 中文 PDF 无法渲染，直接报错便于发现
    if font_path is None:
        raise RuntimeError("未找到可用中文字体（Windows Fonts 下 simhei/msyh/simsun）")
    font_name = "CJK"
    # TTC 是字体集合，需指定 subfontIndex=0 取第一个子字体；TTF 不需要该参数
    ttf_kwargs = {"subfontIndex": 0} if font_path.suffix.lower() == ".ttc" else {}
    pdfmetrics.registerFont(TTFont(font_name, str(font_path), **ttf_kwargs))

    body_style = ParagraphStyle(
        "cjk_body", fontName=font_name, fontSize=10.5, leading=16, wordWrap="CJK"
    )
    title_style = ParagraphStyle(
        "cjk_title", parent=body_style, fontSize=14, leading=20, spaceAfter=10
    )
    flowables = []
    for line in render_contract(spec).splitlines():
        if not line.strip():
            continue
        # 分支：markdown 标题行 → PDF 标题段落
        if line.startswith("# "):
            flowables.append(Paragraph(line[2:], title_style))
        else:
            flowables.append(Paragraph(_escape_pdf_text(line), body_style))
    path.parent.mkdir(parents=True, exist_ok=True)
    SimpleDocTemplate(str(path), pagesize=A4).build(flowables)


def main() -> None:
    """把 SPECS 全部渲染成 md/docx/pdf 落盘到 data/contracts/，并打印清单。"""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    manifest: list[str] = []
    for spec in SPECS:
        stem = Path(spec.filename).stem
        # 三格式同源输出：正文来自同一份 render_contract，保证内容一致
        (OUTPUT_DIR / spec.filename).write_text(render_contract(spec), encoding="utf-8")
        render_docx(spec, OUTPUT_DIR / "docx" / f"{stem}.docx")
        render_pdf(spec, OUTPUT_DIR / "pdf" / f"{stem}.pdf")
        manifest.append(f"{spec.sample_id}\t{spec.filename}\t{spec.note}")
    print(f"已生成 {len(SPECS)} 份合成合同（md/docx/pdf）-> {OUTPUT_DIR}")
    for line in manifest:
        print(" -", line)


if __name__ == "__main__":
    main()
