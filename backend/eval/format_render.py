"""docx/pdf 合同感排版渲染（生成器 v2）。

与 md 的关系：正文条款文本来自 generate_samples 的正文渲染（同源）——企业样本
render_contract（第X条式）、校服样本 render_uniform_contract（章节式 + 明细表），
本模块只负责把同一份正文按「真实采购合同」的版式排出来：
- A4 页面、宋体正文小四两端对齐 + 首行缩进两字符、条款头黑体、标题居中加粗；
- 首部双方信息表（名称/统一社会信用代码/法定代表人/住所）、签署盖章区、页脚页码；
- markdown 明细表（| 行连续成块）渲染成真表格：docx 用 Word 表格、pdf 用 reportlab Table；
- 付款明细保持文字（一）（二）形式而非表格：PDF 表格文本抽取顺序不可靠，
  金额一致性/预付款规则依赖金额文本可抽，故关键事实字段不进表格（易错点）。

设计约定：所有装饰性内容（信息表/签署区）不引入新的可抽取事实；
金额、日期、条款编号等规则字段只以普通段落出现，保证 parser/extractor 不变。
"""

from __future__ import annotations

import re
import os
import zlib
from pathlib import Path

# 双方法定代表人 / 住所候选池（合成数据，不含真实公司信息）
_LEGAL_REP_POOL = {
    "buyer": ("陈立", "王芳", "刘洋", "周涛"),
    "supplier": ("李强", "张敏", "赵磊", "孙悦"),
}
_ADDRESS_POOL = (
    "浙江省杭州市西湖区文三路 188 号",
    "上海市浦东新区张江路 68 号",
    "北京市海淀区中关村大街 27 号",
    "深圳市南山区科技园南区 12 栋",
)


def _crc(text: str) -> int:
    """稳定伪随机种子：同一 spec 每次重生成输出一致（可复现）。"""
    return zlib.crc32(text.encode("utf-8"))


def _party_facts(spec, side: str) -> dict[str, str]:
    """按 sample_id + 甲方/乙方生成确定性双方信息（名称来自 spec 本身）。

    返回 信用代码/法定代表人/住所 三个字段；同一份样本 docx 与 pdf 共用，
    保证两格式展示一致。全部为合成的占位信息。
    """
    seed = f"{spec.sample_id}:{side}"
    reps = _LEGAL_REP_POOL[side]
    return {
        # 18 位统一社会信用代码：固定前缀 + 稳定散列（仅演示用）
        "credit_code": f"91330106MA{_crc(seed + ':code'):08X}",
        "legal_rep": reps[_crc(seed + ":rep") % len(reps)],
        "address": _ADDRESS_POOL[_crc(seed + ":addr") % len(_ADDRESS_POOL)],
    }


# ---- md 正文 → 渲染块（条款头识别 / markdown 表格）----

# 条款头：第X条（企业/GF）或章节式「一、二、…」（校服/政采示范文本），docx/pdf 一律黑体加粗
_HEADING_RE = re.compile(r"^(?:第[0-9一二三四五六七八九十百千万零〇]+条|[一二三四五六七八九十]+、)")
# markdown 表格行：行首 |（明细表/附件清单用），连续 | 行收成表格块
_TABLE_ROW_RE = re.compile(r"^\s*\|")
# markdown 分隔行单元格：纯 - 或 :---:（| --- | --- |）
_SEP_CELL_RE = re.compile(r"^:?-+:?$")


def _split_md_table_row(line: str) -> list[str]:
    """拆 markdown 表格行 → 单元格文本列表（剥首尾 |、逐格去空白）。"""
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_md_separator(cells: list[str]) -> bool:
    """分隔行判断：整行单元格都是 -/: 组合（| --- | --- |）→ True（渲染时丢弃）。"""
    return bool(cells) and all(_SEP_CELL_RE.fullmatch(cell) for cell in cells)


def parse_md_blocks(lines: list[str]) -> list[tuple]:
    """把 md 正文行解析成渲染块：("heading"|"para", 文本) 或 ("table", 单元格二维列表)。

    规则：正文前的合同头（标题/编号/甲乙方）整段跳过——排版外壳已自绘这些内容，
    防止与 spec.title 等重复出现；正文起始=第一个条款/章节头；连续 | 行收成表格块，
    分隔行丢弃。返回块列表供 docx/pdf 两种渲染器共用同一套正文。
    """
    blocks: list[tuple] = []
    started = False  # 是否已进入正文（首个条款/章节头之后），避免重复渲染合同头
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        # 分支 1：空行 → 跳过
        if not line:
            i += 1
            continue
        # 分支 2：还没到正文起始 → 合同头段落直接跳过
        if not started:
            if not _HEADING_RE.match(line):
                i += 1
                continue
            started = True
        # 分支 3：markdown 表格 → 连续收 | 行成表格块
        if _TABLE_ROW_RE.match(line):
            rows: list[list[str]] = []
            while i < len(lines) and _TABLE_ROW_RE.match(lines[i].strip()):
                cells = _split_md_table_row(lines[i])
                if not _is_md_separator(cells):
                    rows.append(cells)
                i += 1
            blocks.append(("table", rows))
            continue
        # 分支 4：条款/章节头 → heading 块；其余正文 → para 块
        blocks.append(("heading", line) if _HEADING_RE.match(line) else ("para", line))
        i += 1
    return blocks


def _add_md_table(doc, rows: list[list[str]]) -> None:
    """把 markdown 明细表行渲染成 Word 表格（表头行黑体，与双方信息表同风格）。"""
    from docx.enum.table import WD_TABLE_ALIGNMENT

    ncols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=ncols, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ridx, row in enumerate(rows):
        # 表头行黑体加粗、数据行宋体，与首部双方信息表一致
        font, bold = ("黑体", True) if ridx == 0 else ("宋体", False)
        for cidx in range(ncols):
            text = row[cidx] if cidx < len(row) else ""
            _fill_cell(table.rows[ridx].cells[cidx], text, font=font, size=10.5, bold=bold)


def _pdf_md_table(rows: list[list[str]], width: float):
    """把 markdown 明细表行渲染成 reportlab Table（按内容长度加权分列宽）。"""
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph, Table, TableStyle

    fonts = _PDF_FONTS
    head = ParagraphStyle("t_head", fontName=fonts["heading"], fontSize=10.5, leading=14)
    cell = ParagraphStyle("t_cell", fontName=fonts["body"], fontSize=10.5, leading=14)
    ncols = max(len(row) for row in rows)
    data: list[list[Paragraph]] = []
    weights = [1] * ncols  # 列权重：按本列最长文本粗估，长文本列多分宽度
    for ridx, row in enumerate(rows):
        line: list[Paragraph] = []
        for cidx in range(ncols):
            text = row[cidx] if cidx < len(row) else ""
            style = head if ridx == 0 else cell
            line.append(Paragraph(_escape_pdf_text(text), style))
            weights[cidx] = max(weights[cidx], max(1, len(text)))
        data.append(line)
    total = sum(weights)
    table = Table(data, colWidths=[width * w / total for w in weights])
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#9E9E9E")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]
        )
    )
    return table


# ---- docx 渲染 ----


def _style_run(run, font: str = "宋体", size: float = 12, bold: bool = False) -> None:
    """给 run 设字号/加粗，并同时设中文字体（w:eastAsia）与西文 Times New Roman。

    易错点：run.font.name 只写 ascii/hAnsi，中文渲染看 eastAsia，
    不设 eastAsia 会回落到主题字体（等线），观感不像合同。
    """
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), font)


def _add_para(
    doc,
    text: str,
    *,
    font: str = "宋体",
    size: float = 12,
    bold: bool = False,
    align=None,
    indent: bool = False,
    before: float = 0,
    after: float = 0,
) -> None:
    """docx 段落小工厂：统一行距 1.5、可选居中/首行缩进两字符/段前段后距。"""
    from docx.shared import Pt

    p = doc.add_paragraph()
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.line_spacing = 1.5
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    # 首行缩进按字号两字符折算（中文习惯两字符起头）
    if indent:
        pf.first_line_indent = Pt(size * 2)
    _style_run(p.add_run(text), font=font, size=size, bold=bold)


def _fill_cell(cell, text: str, *, font: str = "宋体", size: float = 10.5, bold: bool = False) -> None:
    """表格单元格写字：清空既有 run 再写字（防残留空 run 干扰加粗读取/脏 XML）。"""
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)
    _style_run(paragraph.add_run(text), font=font, size=size, bold=bold)


def _party_table(doc, spec) -> None:
    """首部双方信息表：两列（甲方/乙方）× 三行（信用代码/法定代表人/住所）。"""
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table = doc.add_table(rows=4, cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    facts = {side: _party_facts(spec, side) for side in ("buyer", "supplier")}
    rows = [
        (f"甲方（采购方）：{spec.buyer}", f"乙方（供应商）：{spec.supplier}"),
        (
            f"统一社会信用代码：{facts['buyer']['credit_code']}",
            f"统一社会信用代码：{facts['supplier']['credit_code']}",
        ),
        (
            f"法定代表人：{facts['buyer']['legal_rep']}",
            f"法定代表人：{facts['supplier']['legal_rep']}",
        ),
        (f"住所：{facts['buyer']['address']}", f"住所：{facts['supplier']['address']}"),
    ]
    for ridx, (left, right) in enumerate(rows):
        # 表头行用黑体与正文区分；其余行宋体
        font, bold = ("黑体", True) if ridx == 0 else ("宋体", False)
        _fill_cell(table.rows[ridx].cells[0], left, font=font, bold=bold)
        _fill_cell(table.rows[ridx].cells[1], right, font=font, bold=bold)


def _signature_table(doc, spec) -> None:
    """尾部签署区：无边框两列表，甲方乙方各一行盖章/签字/日期。"""
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    facts = {side: _party_facts(spec, side) for side in ("buyer", "supplier")}
    rows = [
        (f"甲方（盖章）：{spec.buyer}", f"乙方（盖章）：{spec.supplier}"),
        (
            f"法定代表人或授权代表（签字）：{facts['buyer']['legal_rep']}",
            f"法定代表人或授权代表（签字）：{facts['supplier']['legal_rep']}",
        ),
        ("签署日期：      年    月    日", "签署日期：      年    月    日"),
    ]
    for ridx, (left, right) in enumerate(rows):
        # 日期留空（真实合同盖章时手填），避免与正文签署日期重复产生抽取歧义
        _fill_cell(table.rows[ridx].cells[0], left, size=10.5)
        _fill_cell(table.rows[ridx].cells[1], right, size=10.5)


def _add_footer_page_number(doc) -> None:
    """页脚居中「第 X 页 共 Y 页」（用 PAGE/NUMPAGES 域，Word 打开自动更新）。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _field(instruction: str) -> None:
        # 域=三个 run：fldChar begin + instrText + fldChar end
        for attrs, text in (
            ({"w:fldCharType": "begin"}, None),
            (None, instruction),
            ({"w:fldCharType": "end"}, None),
        ):
            run = p.add_run()
            _style_run(run, size=9)
            if attrs:
                fld = OxmlElement("w:fldChar")
                for k, v in attrs.items():
                    fld.set(qn(k), v)
                run._r.append(fld)
            else:
                instr = OxmlElement("w:instrText")
                instr.set(qn("xml:space"), "preserve")
                instr.text = f" {text} "
                run._r.append(instr)

    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 顺序拼装：静态文本与 PAGE/NUMPAGES 域交替（域在 Word 打开时自动计算）
    run = p.add_run("第 ")
    _style_run(run, size=9)
    _field("PAGE")
    run = p.add_run(" 页  共 ")
    _style_run(run, size=9)
    _field("NUMPAGES")
    run = p.add_run(" 页")
    _style_run(run, size=9)


def render_docx(spec, path: str | Path, body: str | None = None) -> None:
    """按 spec 生成「合同感」Word 版：正文与 md 同源（可传入复用），加合同版式外壳。

    body 缺省时按 spec 类型现取正文（企业 render_contract / 校服 render_uniform_contract），
    兼容旧调用方；生成器 main 已算好正文直接传入，避免重复渲染。
    """
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    doc = Document()
    section = doc.sections[0]
    # A4 竖版 + 常规页边距（Word 中文默认），保证与 PDF 观感一致
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

    # Normal 样式：宋体 12pt(小四)、1.5 倍行距、段后 0（合同正文紧凑）
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    # ---- 首部：标题 / 编号 / 甲乙双方 / 信息表 ----
    _add_para(
        doc,
        spec.title,
        font="黑体",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=10,
    )
    _add_para(doc, f"合同编号：{spec.contract_no}", size=10.5, align=WD_ALIGN_PARAGRAPH.RIGHT, after=8)
    _add_para(doc, f"甲方（采购方）：{spec.buyer}")
    _add_para(doc, f"乙方（供应商）：{spec.supplier}", after=6)
    _party_table(doc, spec)

    # ---- 正文条款：与 md 同源；支持第X条/章节式头 + markdown 明细表 ----
    if body is None:
        from backend.eval.generate_samples import (
            TechServiceSampleSpec,
            UniformSampleSpec,
            render_contract,
            render_tech_service_contract,
            render_uniform_contract,
        )

        # 分支：技术开发式/校服式/企业式正文各走各自渲染器（docx/pdf 与 md 同源）
        if isinstance(spec, TechServiceSampleSpec):
            body = render_tech_service_contract(spec)
        elif isinstance(spec, UniformSampleSpec):
            body = render_uniform_contract(spec)
        else:
            body = render_contract(spec)
    for kind, payload in parse_md_blocks(body.splitlines()):
        # 分支 1：条款/章节头 → 黑体加粗、段前留白、顶格（与正文缩进区分层级）
        if kind == "heading":
            _add_para(
                doc,
                payload,
                font="黑体",
                size=12,
                bold=True,
                before=10,
                after=4,
            )
        # 分支 2：markdown 明细表 → 真 Word 表格（明细行数多时自动分页）
        elif kind == "table":
            _add_md_table(doc, payload)
        # 分支 3：条款正文 → 宋体、首行缩进两字符、两端对齐
        else:
            _add_para(doc, payload, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ---- 尾部：无正文提示 + 签署区 ----
    _add_para(doc, "（以下无正文，为签署页）", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=10)
    _signature_table(doc, spec)
    _add_footer_page_number(doc)
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


# ---- PDF 渲染 ----


def _cjk_font_path() -> Path | None:
    """找系统可嵌入中文字体：正文宋体(simsun)优先，标题黑体(simhei)次之。"""
    font_dir = Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts"
    # 分支：simsun.ttc 是集合字体，需 subfontIndex=0；msyh 作宋体缺失时的兜底
    for name in ("simsun.ttc", "msyh.ttc", "simhei.ttf", "simkai.ttf"):
        path = font_dir / name
        if path.exists():
            return path
    return None


_PDF_FONTS: dict[str, str] = {}


def _ensure_pdf_fonts() -> dict[str, str]:
    """注册 PDF 中文字体一次，返回 {body: 字体名, heading: 字体名}。

    宋体/黑体以 TTF/TTC 嵌入；找不到宋体时退回微软雅黑（保证中文可渲染）。
    """
    if _PDF_FONTS:
        return _PDF_FONTS
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_dir = Path(__import__("os").environ.get("WINDIR", "C:/Windows")) / "Fonts"
    candidates = (
        ("Song", "simsun.ttc", 0),  # simsun.ttc 第 0 子字体即宋体
        ("MSYH", "msyh.ttc", 0),
        ("Hei", "simhei.ttf", None),
        ("Kai", "simkai.ttf", None),
    )
    registered: dict[str, str] = {}
    for name, filename, subfont in candidates:
        path = font_dir / filename
        if not path.exists():
            continue
        kwargs = {"subfontIndex": subfont} if subfont is not None else {}
        pdfmetrics.registerFont(TTFont(name, str(path), **kwargs))
        registered[name] = name
    # 分支：找不到任何中文字体 → 明确报错（生成前暴露环境问题）
    if not registered:
        raise RuntimeError("未找到可用中文字体（Windows Fonts 下 simsun/msyh/simhei）")
    body = "Song" if "Song" in registered else ("MSYH" if "MSYH" in registered else next(iter(registered)))
    heading = "Hei" if "Hei" in registered else body
    _PDF_FONTS.update({"body": body, "heading": heading})
    return _PDF_FONTS


def _escape_pdf_text(text: str) -> str:
    """reportlab Paragraph 走 XML，转义 & < > 防解析报错。"""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pdf_party_table(spec, width: float):
    """首部双方信息表（reportlab Table）：grid 细线 + 表头黑体。"""
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph, Table, TableStyle

    fonts = _PDF_FONTS
    facts = {side: _party_facts(spec, side) for side in ("buyer", "supplier")}
    head = ParagraphStyle("phead", fontName=fonts["heading"], fontSize=10.5, leading=15)
    cell = ParagraphStyle("pcell", fontName=fonts["body"], fontSize=10.5, leading=15)
    rows = [
        [Paragraph(_escape_pdf_text(f"甲方（采购方）：{spec.buyer}"), head),
         Paragraph(_escape_pdf_text(f"乙方（供应商）：{spec.supplier}"), head)],
        [Paragraph(_escape_pdf_text(f"统一社会信用代码：{facts['buyer']['credit_code']}"), cell),
         Paragraph(_escape_pdf_text(f"统一社会信用代码：{facts['supplier']['credit_code']}"), cell)],
        [Paragraph(_escape_pdf_text(f"法定代表人：{facts['buyer']['legal_rep']}"), cell),
         Paragraph(_escape_pdf_text(f"法定代表人：{facts['supplier']['legal_rep']}"), cell)],
        [Paragraph(_escape_pdf_text(f"住所：{facts['buyer']['address']}"), cell),
         Paragraph(_escape_pdf_text(f"住所：{facts['supplier']['address']}"), cell)],
    ]
    table = Table(rows, colWidths=[width / 2 - 2, width / 2 - 2])
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#9E9E9E")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def _pdf_signature_table(spec, width: float):
    """尾部签署区（reportlab Table）：无边框，甲乙两列。"""
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph, Table, TableStyle

    fonts = _PDF_FONTS
    facts = {side: _party_facts(spec, side) for side in ("buyer", "supplier")}
    style = ParagraphStyle("sig", fontName=fonts["body"], fontSize=10.5, leading=16)
    rows = [
        [Paragraph(_escape_pdf_text(f"甲方（盖章）：{spec.buyer}"), style),
         Paragraph(_escape_pdf_text(f"乙方（盖章）：{spec.supplier}"), style)],
        [Paragraph(_escape_pdf_text(f"法定代表人或授权代表（签字）：{facts['buyer']['legal_rep']}"), style),
         Paragraph(_escape_pdf_text(f"法定代表人或授权代表（签字）：{facts['supplier']['legal_rep']}"), style)],
        [Paragraph("签署日期：      年    月    日", style), Paragraph("签署日期：      年    月    日", style)],
    ]
    table = Table(rows, colWidths=[width / 2 - 2, width / 2 - 2])
    table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP")]))
    return table


def render_pdf(spec, path: str | Path, body: str | None = None) -> None:
    """按 spec 生成「合同感」PDF：与 docx 同一套版式；body 复用规则同 render_docx。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    fonts = _ensure_pdf_fonts()
    body_font, heading_font = fonts["body"], fonts["heading"]

    # 正文：宋体 12pt(小四) 1.5 倍行距首行缩进两字符两端对齐；条款头黑体顶格
    title = ParagraphStyle("title", fontName=heading_font, fontSize=16, leading=24, alignment=1, spaceAfter=14)
    meta = ParagraphStyle("meta", fontName=body_font, fontSize=10.5, leading=16, alignment=2)
    party = ParagraphStyle("party", fontName=body_font, fontSize=12, leading=20)
    clause_head = ParagraphStyle(
        "clause", fontName=heading_font, fontSize=12, leading=24, spaceBefore=10, spaceAfter=4, keepWithNext=1
    )
    body_style = ParagraphStyle(
        "body",
        fontName=body_font,
        fontSize=12,
        leading=24,
        firstLineIndent=24,
        alignment=4,  # JUSTIFY：中文合同正文两端对齐
    )
    sig_note = ParagraphStyle("sig_note", fontName=body_font, fontSize=10.5, leading=16, alignment=1)

    def _footer(canvas, doc_) -> None:
        """页脚居中「第 N 页」（canvas 绘制，reportlab 分页自动触发）。"""
        canvas.saveState()
        canvas.setFont(body_font, 9)
        canvas.drawCentredString(A4[0] / 2, 36, f"第 {canvas.getPageNumber()} 页")
        canvas.restoreState()

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=72,
        rightMargin=72,
        topMargin=72,
        bottomMargin=64,
        title=spec.title,
        author="ContractAgent 合成样本",
    )
    story = [
        Paragraph(_escape_pdf_text(spec.title), title),
        Paragraph(_escape_pdf_text(f"合同编号：{spec.contract_no}"), meta),
        Spacer(1, 8),
        Paragraph(_escape_pdf_text(f"甲方（采购方）：{spec.buyer}"), party),
        Paragraph(_escape_pdf_text(f"乙方（供应商）：{spec.supplier}"), party),
        Spacer(1, 8),
        _pdf_party_table(spec, A4[0] - 144),
        Spacer(1, 10),
    ]
    # 正文条款：与 docx 同源（支持章节式头与 markdown 明细表，与 docx 渲染规则一致）
    if body is None:
        from backend.eval.generate_samples import (
            TechServiceSampleSpec,
            UniformSampleSpec,
            render_contract,
            render_tech_service_contract,
            render_uniform_contract,
        )

        # 分支：技术开发式/校服式/企业式正文各走各自渲染器（与 render_docx 同规则）
        if isinstance(spec, TechServiceSampleSpec):
            body = render_tech_service_contract(spec)
        elif isinstance(spec, UniformSampleSpec):
            body = render_uniform_contract(spec)
        else:
            body = render_contract(spec)
    for kind, payload in parse_md_blocks(body.splitlines()):
        # 分支 1：条款/章节头 → 黑体、段前留白、不换页断开（keepWithNext）
        if kind == "heading":
            story.append(Paragraph(_escape_pdf_text(payload), clause_head))
        # 分支 2：markdown 明细表 → reportlab Table
        elif kind == "table":
            story.append(_pdf_md_table(payload, A4[0] - 144))
        # 分支 3：正文段落 → 宋体小四两端对齐首行缩进
        else:
            story.append(Paragraph(_escape_pdf_text(payload), body_style))
    story.extend(
        [
            Spacer(1, 10),
            Paragraph("（以下无正文，为签署页）", sig_note),
            Spacer(1, 10),
            _pdf_signature_table(spec, A4[0] - 144),
        ]
    )
    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
